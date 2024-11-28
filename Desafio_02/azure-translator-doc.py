import requests, os
from docx import Document
from dotenv import load_dotenv

load_dotenv()
subscription_key = os.getenv("AZURE_KEY")
endpoint = os.getenv("AZURE_ENDPOINT")
location = os.getenv("SERVER_LOCATION")
language = os.getenv("CHOSEN_LANGUAGE")


def translator_text(text, target_language):
    path = "/translate"
    constructed_url = endpoint + path

    headers = {
        "Ocp-Apim-Subscription-Key": subscription_key,
        "Ocp-Apim-Subscription-Region": location,
        "Content-type": "application/json",
        "X-ClientTraceId": str(os.urandom(16))
    }

    body = [
        { "text": text }
    ]

    params = {
        "api-version": "3.0",
        "from": "en",
        "to": [ "pt-br" ]
    }

    request = requests.post(
        constructed_url, 
        params=params, 
        headers=headers, 
        json=body
    )

    response = request.json()

    return response[0]["translations"][0]["text"]


def translator_document(path):
    document = Document(path)
    full_text = []

    for paragraph in document.paragraphs:
        translated_text = translator_text(paragraph.text, language)
        full_text.append(translated_text)

    translated_doc = Document()

    for line in full_text:
        translated_doc.add_paragraph(line)

    path_translated = path.replace(".docx", f"_{ language }.docx")
    translated_doc.save(path_translated)

    return path_translated


input_file = "example.docx"
translator_document(input_file)